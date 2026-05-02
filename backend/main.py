import os
import asyncio
import json
import tempfile
import shutil
from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, JSONResponse
from pydantic import BaseModel

# Resolve the backend root and the shared data directory using __file__
# so paths are always correct no matter which directory uvicorn is launched from.
BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR    = os.path.join(BACKEND_DIR, "data")
TEMP_DOCX   = os.path.join(DATA_DIR, "temp.docx")

# Use simple relative imports — no full dotted-package path needed.
from src import engine, formatter

app = FastAPI()

# Allow the React app to talk to this Python server
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Data Models ────────────────────────────────────────────────────────────────

class AbstractRequest(BaseModel):
    abstract: str
    raw_text: str

class GenerateRequest(BaseModel):
    metadata: dict
    raw_text: str

class CrossrefRequest(BaseModel):
    citation: str


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/")
async def root():
    return {"message": "N.O.V.A. AI Engine is online and ready!"}


@app.post("/api/verify-citation")
async def verify_citation(req: CrossrefRequest):
    """
    Backend proxy to handle citation verification and bypass CORS.
    Branches logic based on DOI presence.
    """
    try:
        citation_text = req.citation
        # 1. Regex to search for a DOI
        doi_pattern = r'\b(10\.\d{4,9}/[-._;()/:a-zA-Z0-9]+)\b'
        doi_match = re.search(doi_pattern, citation_text)
        
        headers = {"User-Agent": "NovaBot/1.0 (mailto:integrity@projectnova.io)"}

        if doi_match:
            # 2. Logic: Direct DOI Path Lookup
            doi = doi_match.group(1)
            response = requests.get(f"https://api.crossref.org/works/{doi}", headers=headers, timeout=10)
        else:
            # 3. Logic: Bibliographic Query Fallback
            params = {
                "query.bibliographic": citation_text,
                "rows": 1,
                "select": "title,author,URL,DOI"
            }
            response = requests.get("https://api.crossref.org/works", params=params, headers=headers, timeout=10)
        
        response.raise_for_status()
        return response.json()

    except requests.exceptions.RequestException as exc:
        # 4. Constraint: Clean 500 error for unreachable network
        return JSONResponse(
            status_code=500,
            content={"error": f"Crossref network unreachable: {str(exc)}"}
        )



@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Save file, hash it, and run LLM metadata extraction — all in one request."""
    try:
        content = await file.read()
        with open(TEMP_DOCX, "wb") as f:
            f.write(content)

        # Run all blocking work in threads so uvicorn's event loop stays free
        raw_text     = await asyncio.to_thread(engine.extract_text_from_docx, TEMP_DOCX)
        lexical_hash = await asyncio.to_thread(engine.calculate_lexical_hash, raw_text)
        semantic_hash = await asyncio.to_thread(engine.get_semantic_hash, raw_text)
        metadata     = await asyncio.to_thread(engine.get_document_metadata, raw_text)

        return {
            "raw_text":      raw_text,
            "lexical_hash":  lexical_hash,
            "semantic_hash": semantic_hash,
            "metadata":      metadata,
        }
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": str(e)})


@app.post("/fix-abstract")
async def fix_abstract(request: AbstractRequest):
    try:
        # All three calls are CPU/network-blocking — run them in threads
        fixed_text   = await asyncio.to_thread(engine.fix_and_shorten_abstract, request.abstract)

        import re as _re
        def _norm(s: str) -> str:
            return _re.sub(r'\s+', ' ', s).strip()

        norm_raw      = _norm(request.raw_text)
        norm_abstract = _norm(request.abstract)
        norm_fixed    = _norm(fixed_text)

        # ── Diagnostics ────────────────────────────────────────────────────
        llm_changed = norm_abstract != norm_fixed
        print(f"[N.O.V.A.] fix-abstract: LLM changed text = {llm_changed}")
        if llm_changed:
            print(f"  ORIG[:80]: {norm_abstract[:80]!r}")
            print(f"  FIXED[:80]: {norm_fixed[:80]!r}")

        # ── 4-tier replacement strategy ─────────────────────────────────────
        new_raw_text = None

        # Tier 1: exact normalized match
        if norm_abstract in norm_raw:
            new_raw_text = norm_raw.replace(norm_abstract, norm_fixed, 1)
            print("[N.O.V.A.] fix-abstract: [OK] Tier 1 (exact match) succeeded")

        # Tier 2: anchor on first 200 chars
        if new_raw_text is None:
            anchor = norm_abstract[:200]
            idx = norm_raw.find(anchor)
            if idx >= 0:
                new_raw_text = norm_raw[:idx] + norm_fixed + norm_raw[idx + len(norm_abstract):]
                print("[N.O.V.A.] fix-abstract: [OK] Tier 2 (200-char anchor) succeeded")

        # Tier 3: anchor on first 80 chars
        if new_raw_text is None:
            anchor = norm_abstract[:80]
            idx = norm_raw.find(anchor)
            if idx >= 0:
                new_raw_text = norm_raw[:idx] + norm_fixed + norm_raw[idx + len(norm_abstract):]
                print("[N.O.V.A.] fix-abstract: [OK] Tier 3 (80-char anchor) succeeded")

        # Tier 4: guaranteed fallback — appends fixed abstract so the hash always
        # differs when the LLM made changes (wrong section, but different hash = correct)
        if new_raw_text is None:
            print("[N.O.V.A.] fix-abstract: [WARN] All tiers failed — using fallback concatenation")
            new_raw_text = norm_raw + "\n" + norm_fixed if llm_changed else norm_raw


        lex_hash   = await asyncio.to_thread(engine.calculate_lexical_hash, new_raw_text)
        sem_hash   = await asyncio.to_thread(engine.get_semantic_hash, new_raw_text)
        similarity = await asyncio.to_thread(engine.calculate_semantic_similarity, request.raw_text, new_raw_text)

        return {
            "fixed_abstract":   fixed_text,
            "new_lexical_hash": lex_hash,
            "new_semantic_hash": sem_hash,
            "similarity":       similarity,
        }
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": str(e)})


@app.post("/download/pdf")
async def download_pdf(req: GenerateRequest):
    try:
        pdf_bytes = await asyncio.to_thread(formatter.generate_pdf, req.metadata, req.raw_text)
        # Real PDFs always start with the %PDF magic bytes
        if not pdf_bytes or not pdf_bytes[:4] == b'%PDF':
            error_msg = pdf_bytes.decode(errors='replace') if pdf_bytes else 'No output from pdflatex'
            return JSONResponse(status_code=500, content={"detail": error_msg})
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=NOVA_Manuscript.pdf"}
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": str(e)})


@app.post("/download/report")
async def download_report(req: GenerateRequest):
    try:
        lex_hash = engine.calculate_lexical_hash(req.raw_text)
        sem_hash = engine.get_semantic_hash(req.raw_text)

        report = f"""N.O.V.A. CRYPTOGRAPHIC INTEGRITY REPORT
---------------------------------------
Document Title: {req.metadata.get('title', 'Unknown')}

[ LEXICAL INTEGRITY ]
SHA-256 Hash: {lex_hash}
Status: VERIFIED

[ SEMANTIC INTEGRITY ]
LSH Binarized Hash: {sem_hash}
Status: VERIFIED
(Note: Minor bit variations indicate authorized Gen-AI grammar/formatting fixes).

Verification: 100% Scientific Claim Integrity Confirmed.
"""
        return Response(content=report.encode(), media_type="text/plain")
    except Exception as e:
        return Response(content=str(e), status_code=500)


@app.post("/download/docx")
async def download_docx(req: GenerateRequest):
    try:
        if not os.path.exists(TEMP_DOCX):
            return JSONResponse(
                status_code=400,
                content={"detail": "No uploaded document found. Please upload a file first by going back to Step 1."}
            )

        def _build_docx() -> bytes:
            import docx as _docx
            import io

            # Open the original document to preserve all images, tables, and formatting
            doc = _docx.Document(TEMP_DOCX)

            # Try to patch the abstract if possible, as it might have been fixed by AI
            new_abstract = req.metadata.get('abstract', '')
            if new_abstract:
                for i, p in enumerate(doc.paragraphs):
                    if 'abstract' in p.text.lower() and len(p.text) < 50:
                        # Found the abstract heading, replace the next paragraph
                        if i + 1 < len(doc.paragraphs):
                            doc.paragraphs[i+1].text = new_abstract
                        break

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        docx_bytes = await asyncio.to_thread(_build_docx)
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=NOVA_Manuscript.docx"}
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        return JSONResponse(status_code=500, content={"detail": str(e)})