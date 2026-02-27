import os
import pytest
import docx
from src.engine import extract_text_from_docx

def create_dummy_docx(filepath):
    doc = docx.Document()
    doc.add_heading('Test Document Title', 0)
    doc.add_heading('Introduction', 1)
    doc.add_paragraph('This is a test paragraph.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Data 1'
    table.cell(1, 1).text = 'Data 2'
    
    doc.save(filepath)

def test_extract_text_from_docx_uses_unstructured(tmp_path):
    test_file = tmp_path / "test.docx"
    create_dummy_docx(str(test_file))
    
    text = extract_text_from_docx(str(test_file))
    
    # We expect unstructured to parse it properly including the table as HTML or structured text
    assert 'Test Document Title' in text
    assert 'Introduction' in text
    assert 'This is a test paragraph.' in text
    # Ideally, we mapped headings or tables correctly
    assert 'Header 1' in text
    assert 'Data 1' in text
