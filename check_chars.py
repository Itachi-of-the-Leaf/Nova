import docx

doc = docx.Document("assets/Sample 3.docx")
chars = set()
for para in doc.paragraphs:
    for c in para.text:
        if ord(c) > 255:
            chars.add(c)
            
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for c in cell.text:
                if ord(c) > 255:
                    chars.add(c)

print("Non-ASCII chars in Sample 3:")
for c in sorted(chars):
    print(f"{c} (U+{ord(c):04X})")
