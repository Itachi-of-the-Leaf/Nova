import docx
doc = docx.Document("assets/Sample_3_Input.docx")
print("Sample 3 Tables:", len(doc.tables))

docX = docx.Document("assets/Sample_X_Input.docx")
print("Sample X Tables:", len(docX.tables))
